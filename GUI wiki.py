# Guiwiki.py
import wikipedia

#Python to docx
from docx import Document
def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    #Summary สำหรับบทความที่สรุป
    data = wikipedia.summary(keyword)
    
    #Page + Content = บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content
    
    doc = Document() #ส่งข้อมูลที่ได้ เข้า Word
    doc.add_heading(keyword,0)
    
    doc.add_paragraph(data2)
    doc.save(keyword + ".docx")
    print("สร้างไฟล์สำเร็จ")

#Change Language to Thai
wikipedia.set_lang("th")


from tkinter import * 
from tkinter import ttk
from tkinter import messagebox

GUI = Tk()
GUI.title('โปรแกรม Wiki')
GUI.geometry('400x300')

#ขนาดตัวอักษร
Font1 = ('Angsana New',15)

#ผลลัพธ์

L1 = ttk.Label(GUI,text='ค้นหาบทความ',font=Font1)
L1.pack()

#ช่องค้นหาข้อมูล
v_search = StringVar()
E1 = ttk.Entry(GUI,textvariable = v_search,font=Font1,width=30)
E1.pack(pady=20)

#Search Button
def Search():
    keyword = v_search.get()#.get ใช้ได้แค่ StringVar เท่านั้น
    try: #Try = ลองค้นหาว่าได้ผลลัพท์หรือไม่ หากได้ ให้ไป Step ต่อไป
        language = v_radio.get() #th / en / zh
        Wiki(keyword,language)
        messagebox.showinfo('Record sucess','Research Success save by Word')
    except: # หาก Run แล้วเจอปัญหา แสดงข้อความแจ้งเตือน
        messagebox.showwarning('Keyword Error','Please check your keyword again!')

    # print(wikipedia.search(keyword))
    # result = wikipedia.summary(keyword)
    # print(result)

    
B1 = ttk.Button(GUI,text="Search",command=Search)
B1.pack(ipadx=15,ipady=10) #ipad x,y คือการขยายขนาดของปุ่มภายในตามแกน X,Y

# Select language
F1 = Frame(GUI)
F1.pack(pady=20)

v_radio = StringVar() #ช่องเก็บข้อมูลภาษา

RB1 = ttk.Radiobutton(F1,text='Thai',variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1,text='English',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='Chinese',variable=v_radio,value='zh')
RB1.invoke() #invoke = สั่งให้ค่าเริ่มต้อเป็นภาษาไทย

RB1.grid(row=0,column=0)  
RB2.grid(row=0,column=1) 
RB3.grid(row=0,column=2) 


GUI.mainloop()
