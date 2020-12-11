#Test-docx.py

from docx import Document 
import wikipedia

def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    #Summary สำหรับบทความที่สรุป
    data = wikipedia.summary(keyword)
    
    #Page + Content = บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content
    
    doc = Document() #ส่งข้อมูลที่ได้ เข้า Word
    doc.add_heading(keyword,0)
    
    doc.add_paragraph(data2)
    doc.save(keyword + ".docx")
    print("สร้างไฟล์ สำเร็จ")
try:
    Wiki('asdqweasd','en')
except :
    print('Please check your keyword again.')
